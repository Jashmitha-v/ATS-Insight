"""Extract raw text from uploaded PDF / DOCX resume files."""
import io

import pdfplumber
from docx import Document
from PyPDF2 import PdfReader

from utils.validators import ValidationError


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception:
        text_parts = []

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # Fallback to PyPDF2 if pdfplumber yields nothing (e.g. unusual encoding)
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        fallback_parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(fallback_parts).strip()
    except Exception as exc:
        raise ValidationError(
            "Could not read the PDF file. It may be corrupted or scanned as an image.",
            code="UNREADABLE_PDF",
        ) from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValidationError(
            "Could not read the DOCX file. It may be corrupted.",
            code="UNREADABLE_DOCX",
        ) from exc

    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(p for p in paragraphs if p.strip())


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValidationError("Unsupported file format.", code="UNSUPPORTED_FORMAT")

    if not text or not text.strip():
        raise ValidationError(
            "No readable text was found in the resume. It may be a scanned image "
            "or an empty document.",
            code="NO_TEXT_FOUND",
        )
    return text
